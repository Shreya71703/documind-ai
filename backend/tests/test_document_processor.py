import os
import uuid
import pytest
from unittest.mock import MagicMock, patch
from datetime import datetime, timezone
import docx

from app.core.config import settings
from app.services.document_processor import (
    normalize_text,
    process_document,
    DocumentProcessingError,
    UnsupportedFileTypeError,
    StoredFileMissingError,
    EmptyExtractedContentError,
    ChunkResult,
    DocumentProcessingResult
)

# -------------------------------------------------------------
# Unit Tests for Text Normalization
# -------------------------------------------------------------

def test_normalize_text():
    # Test trailing/leading spaces and line endings
    raw = "  Hello \r\n World  "
    assert normalize_text(raw) == "Hello \n World"

    # Test collapsing multiple spaces/tabs
    raw = "Hello\t\t   World"
    assert normalize_text(raw) == "Hello World"

    # Test collapsing three or more newlines to exactly two
    raw = "Line 1\n\n\n\nLine 2"
    assert normalize_text(raw) == "Line 1\n\nLine 2"

    # Empty/None checks
    assert normalize_text("") == ""
    assert normalize_text(None) == ""

# -------------------------------------------------------------
# Unit Tests for Document Processing Service
# -------------------------------------------------------------

def test_process_txt_success(tmp_path):
    doc_id = uuid.uuid4()
    file_path = tmp_path / "test.txt"
    # Create text longer than CHUNK_SIZE (1000) to verify split
    paragraph = "This is a long sentence repeated to exceed chunk size. " * 30
    file_path.write_text(paragraph, encoding="utf-8")

    result = process_document(doc_id, str(file_path), "test.txt")

    assert isinstance(result, DocumentProcessingResult)
    assert result.document_id == doc_id
    assert result.extracted_text_length == len(paragraph)
    assert result.chunk_count > 1
    assert result.status == "ready"

    # Verify chunk structure and metadata
    for idx, chunk in enumerate(result.chunks):
        assert isinstance(chunk, ChunkResult)
        assert len(chunk.content) <= settings.CHUNK_SIZE
        assert chunk.metadata["chunk_index"] == idx
        assert chunk.metadata["document_id"] == str(doc_id)
        assert chunk.metadata["source_filename"] == "test.txt"
        assert chunk.metadata["file_type"] == "txt"

def test_process_txt_utf8_bom(tmp_path):
    doc_id = uuid.uuid4()
    file_path = tmp_path / "bom.txt"
    content = "Hello from UTF-8 BOM"
    # Write with BOM encoding
    file_path.write_text(content, encoding="utf-8-sig")

    result = process_document(doc_id, str(file_path), "bom.txt")
    assert result.extracted_text_length == len(content)
    assert result.chunks[0].content == content

def test_process_md_success(tmp_path):
    doc_id = uuid.uuid4()
    file_path = tmp_path / "test.md"
    content = "# Markdown Heading\n\nParagraph text here."
    file_path.write_text(content, encoding="utf-8")

    result = process_document(doc_id, str(file_path), "test.md")
    assert result.chunk_count == 1
    assert result.chunks[0].content == "# Markdown Heading\n\nParagraph text here."
    assert result.chunks[0].metadata["file_type"] == "md"

def test_process_docx_success(tmp_path):
    doc_id = uuid.uuid4()
    file_path = tmp_path / "test.docx"

    # Create real DOCX file
    doc = docx.Document()
    doc.add_paragraph("Paragraph 1 of Word document.")
    doc.add_paragraph("Paragraph 2 of Word document.")
    doc.save(str(file_path))

    result = process_document(doc_id, str(file_path), "test.docx")
    assert result.status == "ready"
    assert result.chunk_count == 1
    # Check paragraphs are joined with double newline
    expected_content = "Paragraph 1 of Word document.\n\nParagraph 2 of Word document."
    assert result.chunks[0].content == expected_content
    assert result.chunks[0].metadata["file_type"] == "docx"

@patch("app.services.document_processor.PdfReader")
def test_process_pdf_success(mock_pdf_reader_class, tmp_path):
    doc_id = uuid.uuid4()
    file_path = tmp_path / "test.pdf"
    file_path.write_text("dummy pdf content", encoding="utf-8") # Must exist on disk

    # Mock PdfReader behavior
    mock_page1 = MagicMock()
    mock_page1.extract_text.return_value = "Text on Page 1."
    mock_page2 = MagicMock()
    mock_page2.extract_text.return_value = "Text on Page 2."

    mock_reader = MagicMock()
    mock_reader.pages = [mock_page1, mock_page2]
    mock_pdf_reader_class.return_value = mock_reader

    result = process_document(doc_id, str(file_path), "test.pdf")
    
    assert result.status == "ready"
    assert result.chunk_count == 2
    assert result.chunks[0].content == "Text on Page 1."
    assert result.chunks[0].metadata["page_number"] == 1
    assert result.chunks[1].content == "Text on Page 2."
    assert result.chunks[1].metadata["page_number"] == 2
    assert result.chunks[0].metadata["file_type"] == "pdf"

@patch("app.services.document_processor.PdfReader")
def test_process_pdf_empty_text(mock_pdf_reader_class, tmp_path):
    doc_id = uuid.uuid4()
    file_path = tmp_path / "empty.pdf"
    file_path.write_text("dummy", encoding="utf-8")

    mock_page = MagicMock()
    mock_page.extract_text.return_value = "   " # Whitespace only

    mock_reader = MagicMock()
    mock_reader.pages = [mock_page]
    mock_pdf_reader_class.return_value = mock_reader

    with pytest.raises(EmptyExtractedContentError) as exc:
        process_document(doc_id, str(file_path), "empty.pdf")
    assert "no extractable text" in str(exc.value).lower()

@patch("app.services.document_processor.PdfReader")
def test_process_pdf_no_pages(mock_pdf_reader_class, tmp_path):
    doc_id = uuid.uuid4()
    file_path = tmp_path / "nopages.pdf"
    file_path.write_text("dummy", encoding="utf-8")

    mock_reader = MagicMock()
    mock_reader.pages = []
    mock_pdf_reader_class.return_value = mock_reader

    with pytest.raises(EmptyExtractedContentError) as exc:
        process_document(doc_id, str(file_path), "nopages.pdf")
    assert "no pages" in str(exc.value).lower()

def test_process_unsupported_type(tmp_path):
    doc_id = uuid.uuid4()
    file_path = tmp_path / "test.exe"
    file_path.write_text("binary data", encoding="utf-8")

    with pytest.raises(UnsupportedFileTypeError) as exc:
        process_document(doc_id, str(file_path), "test.exe")
    assert "unsupported" in str(exc.value).lower()

def test_process_missing_file():
    doc_id = uuid.uuid4()
    with pytest.raises(StoredFileMissingError) as exc:
        process_document(doc_id, "nonexistent_file.txt", "nonexistent_file.txt")
    assert "file not found" in str(exc.value).lower()

def test_process_empty_text_file(tmp_path):
    doc_id = uuid.uuid4()
    file_path = tmp_path / "empty.txt"
    file_path.write_text("   \n  \t", encoding="utf-8") # Empty space

    with pytest.raises(EmptyExtractedContentError) as exc:
        process_document(doc_id, str(file_path), "empty.txt")
    assert "no extractable text" in str(exc.value).lower()
